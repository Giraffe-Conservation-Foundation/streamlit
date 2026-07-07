"""
Unit Performance Report Dashboard

Pulls live GPS/GSM tracking-device history from EarthRanger for the three
tail/ear tag programmes (SpoorTrack, Savannah Tracking, GSat Solar), computes
the same deployment / battery / fix-rate / performance metrics as the
ER_performReport R scripts, and packages the result into a downloadable
.docx report styled after the existing GCF_unitPerformance Google Docs.
"""

import io
import json
import random
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

sys.path.append(str(Path(__file__).parent.parent))
from shared.auth import require_earthranger_login

ER_SERVER = "https://twiga.pamdas.org"
UNIT_UPDATE_EVENT_TYPE = "7bb99e0c-9d37-405b-b8e7-edca8e9b5d6b"
GCF_LOGO = Path(__file__).parent.parent / "shared" / "logo.png"

# ── Manufacturer configuration ──────────────────────────────────────────────
MANUFACTURERS = {
    "SpoorTrack (tail)": dict(
        provider="spoortrack",
        battery_unit="voltage",
        battery_range=(3.60, 4.20),
        battery_floor=3.6,
        exclude_tagid_contains="EAR",
        since_filter=None,
    ),
    "Savannah Tracking (tail)": dict(
        provider="savannah_tracking_provider",
        battery_unit="voltage",
        battery_range=(3.60, 4.20),
        battery_floor=3.6,
        exclude_tagid_contains=None,
        since_filter=date(2024, 1, 1),
    ),
    "GSat Solar (ear)": dict(
        provider="gsatsolar",
        battery_unit="percentage",
        battery_range=(10, 100),
        battery_floor=10,
        exclude_tagid_contains=None,
        since_filter=None,
    ),
}

# Known manual GSat schedule overrides (serial numbers = collar_key), ported
# from ER_performReport_clean.R since these units' true schedules don't match
# the auto-detected mode-based heuristic.
GSAT_MANUAL_6PERDAY = {
    "009478702", "017230160", "017231465", "017233065", "017249061",
    "017251208", "017252461", "017252768", "017252966", "017256587",
    "017259003", "017260266", "017260282", "017265422", "017267147",
    "017269887", "017270349", "017270364", "033633264", "033635525",
}
GSAT_MANUAL_12PERDAY = {
    "033640921", "033646183", "036003622", "036009884", "017232208",
    "034641506", "033559022", "033592304", "033551805", "033462946",
    "033641069", "033639360", "033464827", "017268806",
}
GSAT_SCHEDULE_CHANGE_UNITS = {"034654541", "036025302"}
GSAT_SCHEDULE_CHANGE_DATE = date(2025, 7, 14)

# Known SpoorTrack test periods (tagID -> ordered list of (start, end, fixes/day),
# first match wins), ported from ER_performReport_clean.R.
SPOORTRACK_TEST_PERIODS = {
    "TAIL-ST1766": [
        (date(2025, 7, 5), date(2025, 9, 6), 98),
        (date(2025, 6, 14), date(2025, 7, 4), 48),
    ],
    "TAIL-ST1759": [
        (date(2025, 7, 5), date(2025, 9, 3), 96),
        (date(2025, 6, 14), date(2025, 9, 4), 48),
    ],
}

EXCLUDE_COUNTRIES = {"AF", "TwigaTracker", "WildScapeVet", "Donor", "Movebank", "White"}
EXCLUDE_GROUP_IDS = {"bbe1967b-6a46-4a58-aca1-5e3f791933dd"}

# The whole report is restricted to this species — some SpoorTrack/Ceres-type devices
# are also deployed on other species (e.g. elephant) for unrelated research programmes.
GIRAFFE_SUBJECT_SUBTYPE = "giraffe"


def main():
    return _main_implementation()


# ═══════════════════════════════════════════════════════════════════════════
# Data fetching
# ═══════════════════════════════════════════════════════════════════════════

@st.cache_data(ttl=3600, show_spinner=False)
def fetch_manufacturer_sources(_er, provider):
    try:
        sources = _er.get_sources(provider=provider)
    except Exception:
        return pd.DataFrame()
    if sources is None or sources.empty:
        return pd.DataFrame()
    # Don't trust the server-side provider= filter alone — re-filter client-side too.
    # (unit_check_dashboard's own get_all_sources() deliberately fetches everything
    # unfiltered and filters by the 'provider' column client-side, which is a strong
    # signal the server-side filter isn't reliable on this EarthRanger instance.)
    if "provider" in sources.columns:
        sources = sources[sources["provider"] == provider]
    sources = sources[sources.get("source_type") == "tracking-device"].copy()
    if sources.empty:
        return sources
    sources["tagID"] = sources["collar_key"].fillna(sources["id"]).astype(str)
    return sources.reset_index(drop=True)


def _extract_ref_id(ref):
    """Extract a bare UUID from an EarthRanger relation field, which — depending on
    server/serializer config — may come through as a nested dict ({'id': ...}), a
    bare UUID string, or a hyperlinked API URL (".../subject/<uuid>/")."""
    if isinstance(ref, dict):
        rid = ref.get("id")
        return str(rid) if rid else None
    if isinstance(ref, str):
        s = ref.rstrip("/")
        return s.rsplit("/", 1)[-1] if "/" in s else s
    return None


@st.cache_data(ttl=1800, show_spinner=False)
def fetch_assignments(_er, source_ids_tuple):
    """One row per subjectsource assignment: source_id, subject_id, subName, depStart, depEnd."""
    if not source_ids_tuple:
        return pd.DataFrame(columns=["source_id", "subject_id", "subName", "depStart", "depEnd"])
    try:
        raw = _er.get_subjectsources(sources=",".join(source_ids_tuple))
    except Exception:
        return pd.DataFrame(columns=["source_id", "subject_id", "subName", "depStart", "depEnd"])
    if raw is None or raw.empty:
        return pd.DataFrame(columns=["source_id", "subject_id", "subName", "depStart", "depEnd"])

    rows = []
    for _, r in raw.iterrows():
        ar = r.get("assigned_range") or {}
        if isinstance(ar, str):
            try:
                ar = json.loads(ar)
            except Exception:
                ar = {}
        if not isinstance(ar, dict):
            ar = {}
        start_raw = ar.get("lower") or ar.get("start_time")
        end_raw = ar.get("upper") or ar.get("end_time")
        dep_start = pd.to_datetime(start_raw, utc=True, errors="coerce")
        dep_end = pd.to_datetime(end_raw, utc=True, errors="coerce")

        subj = r.get("subject")
        subject_id = _extract_ref_id(subj)
        # Best-effort name if the API happened to expand the relation to a dict —
        # the authoritative name lookup is fetch_subject_metadata(), used downstream.
        subject_name = subj.get("name", "") if isinstance(subj, dict) else ""

        src_id = _extract_ref_id(r.get("source"))

        rows.append({
            "source_id": src_id,
            "subject_id": subject_id,
            "subName": subject_name,
            "depStart": dep_start,
            "depEnd": dep_end,
        })
    df = pd.DataFrame(rows).dropna(subset=["source_id", "depStart"])
    return df.reset_index(drop=True)


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_subject_metadata(_er):
    """subject_id -> {'name':.., 'subtype': lowercased}, fetched directly from the
    subjects/ endpoint since the subjectsources relation doesn't reliably expand to
    a named object. subtype is used to restrict the whole report to giraffe subjects
    only — some tracking devices (SpoorTrack/Ceres etc.) are also deployed on other
    species (e.g. elephant) for unrelated research, and those must not leak in here."""
    try:
        subjects = _er.get_subjects(include_inactive=True)
    except Exception:
        return {}
    if subjects is None or subjects.empty or "id" not in subjects.columns:
        return {}
    name_col = subjects["name"] if "name" in subjects.columns else pd.Series([""] * len(subjects), index=subjects.index)
    if "subject_subtype" in subjects.columns:
        subtype_col = subjects["subject_subtype"].astype(str).str.lower()
    else:
        subtype_col = pd.Series([""] * len(subjects), index=subjects.index)
    return {
        str(sid): {"name": name or "", "subtype": subtype}
        for sid, name, subtype in zip(subjects["id"].astype(str), name_col.fillna(""), subtype_col.fillna(""))
    }


def _get_with_retry(_er, path, params=None, max_attempts=3, backoff_seconds=2):
    """Wrap er._get() with retries for connection-level failures.

    erclient's own retry loop only re-attempts on a bad HTTP status code — it
    doesn't catch connection-level exceptions (dropped/reset connections,
    timeouts) at all, so a single transient network blip fails the whole
    request with zero retries. This adds that missing layer.
    """
    last_exc = None
    for attempt in range(1, max_attempts + 1):
        try:
            return _er._get(path, params=params), attempt
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as e:
            last_exc = e
            if attempt < max_attempts:
                time.sleep(backoff_seconds * attempt)
    raise last_exc


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_subject_country_region(_er):
    """subject_id -> (country, region, group_name, group_id), derived from subject group
    names ('COUNTRY_REGION_SPECIES'). group_name/group_id are carried through so failure
    listings can reference the site/region group a unit belonged to for context.

    Returns (mapping, debug_info) — debug_info is a short string describing what
    the raw API response looked like, so an empty result can be diagnosed without
    guessing (this endpoint has no ecoscope wrapper, so we hit it directly)."""
    # requests serializes Python bools as "True"/"False" (capitalised) in the query
    # string; Django REST's boolean filter parsing expects lowercase, so pass strings
    # explicitly rather than relying on Python bool -> str conversion.
    params = {"flat": "true", "include_inactive": "true", "include_hidden": "true"}
    groups = []
    next_url = "subjectgroups/"
    pages_fetched = 0
    retries_used = 0
    first_raw_type = None
    try:
        while next_url and pages_fetched < 20:  # safety cap against runaway pagination
            raw, attempts = _get_with_retry(_er, next_url, params=params if pages_fetched == 0 else None)
            retries_used += attempts - 1
            if pages_fetched == 0:
                first_raw_type = type(raw).__name__
            pages_fetched += 1
            # ecoscope's client already unwraps a top-level {"data": ...} envelope, so
            # what we get back here is either a bare list, or (if paginated) a dict
            # shaped like {"count":.., "next": <url|null>, "results": [...]}.
            if isinstance(raw, dict):
                page_items = raw.get("results", [])
                groups.extend(page_items if isinstance(page_items, list) else [])
                next_url = raw.get("next")
            elif isinstance(raw, list):
                groups.extend(raw)
                next_url = None
            else:
                next_url = None
    except Exception as e:
        return {}, f"subjectgroups/ request raised after {pages_fetched} page(s) ({retries_used} retries used): {e}"

    debug_info = (
        f"raw response type={first_raw_type}, fetched {pages_fetched} page(s) "
        f"({retries_used} retries used), parsed {len(groups)} group(s)"
        + (f", first group keys={list(groups[0].keys())}" if groups and isinstance(groups[0], dict) else "")
    )

    mapping = {}
    for g in groups:
        if not isinstance(g, dict):
            continue
        if g.get("id") in EXCLUDE_GROUP_IDS:
            continue
        gname = g.get("name") or ""
        parts = gname.split("_", 2)
        country = parts[0] if len(parts) > 0 and parts[0] else None
        region = parts[1] if len(parts) > 1 and parts[1] else None
        if country in EXCLUDE_COUNTRIES:
            continue
        for s in (g.get("subjects") or []):
            sid = _extract_ref_id(s)
            if sid and sid not in mapping:
                mapping[sid] = (country, region, gname, g.get("id"))
    debug_info += f", mapped {len(mapping)} subject(s)"
    return mapping, debug_info


def _extract_battery(obs_details, device_status, battery_unit):
    candidates = (
        ["voltage", "battery", "batt", "batt_perc", "bat_soc"]
        if battery_unit == "voltage"
        else ["batt_perc", "battery", "bat_soc", "voltage", "batt"]
    )
    if isinstance(obs_details, dict):
        for f in candidates:
            if f in obs_details and obs_details[f] is not None:
                try:
                    return float(obs_details[f])
                except (TypeError, ValueError):
                    pass
    if isinstance(device_status, list):
        for item in device_status:
            if isinstance(item, dict) and "label" in item and "value" in item:
                label = str(item["label"]).lower()
                if any(t in label for t in ["voltage", "battery", "batt", "bat_soc"]):
                    try:
                        return float(item["value"])
                    except (TypeError, ValueError):
                        pass
    return None


def _fetch_one_source_history(session, headers, source_id, since_iso, until_iso):
    records = []
    url = (
        f"{ER_SERVER}/api/v1.0/observations/"
        f"?source_id={source_id}&since={since_iso}&until={until_iso}"
        f"&include_details=true&page_size=4000"
    )
    while url:
        resp = session.get(url, headers=headers)
        resp.raise_for_status()
        data = resp.json()
        inner = data.get("data", data) if isinstance(data, dict) else data
        if isinstance(inner, dict):
            items = inner.get("results", [])
            url = inner.get("next")
        elif isinstance(inner, list):
            items = inner
            url = None
        else:
            items, url = [], None
        for it in items:
            records.append({
                "source_id": source_id,
                "obsDatetime": it.get("recorded_at"),
                "obs_details": it.get("observation_details"),
                "device_status": it.get("device_status_properties"),
            })
    return records


@st.cache_data(ttl=1800, show_spinner=False)
def fetch_observation_history(_er, source_ids_tuple, since_iso, until_iso, battery_unit):
    headers = _er.auth_headers()
    all_records = []
    with ThreadPoolExecutor(max_workers=10) as pool:
        futures = {
            pool.submit(_fetch_one_source_history, _er._http_session, headers, sid, since_iso, until_iso): sid
            for sid in source_ids_tuple
        }
        for fut in as_completed(futures):
            try:
                all_records.extend(fut.result())
            except Exception:
                pass
    if not all_records:
        return pd.DataFrame(columns=["source_id", "obsDatetime", "battery"])
    df = pd.DataFrame(all_records)
    df["obsDatetime"] = pd.to_datetime(df["obsDatetime"], utc=True, errors="coerce")
    df["battery"] = df.apply(lambda r: _extract_battery(r["obs_details"], r["device_status"], battery_unit), axis=1)
    df = df.dropna(subset=["obsDatetime"])
    return df[["source_id", "obsDatetime", "battery"]].reset_index(drop=True)


@st.cache_data(ttl=1800, show_spinner=False)
def fetch_deactivation_notes(_er, source_ids_tuple):
    """Best-effort end-cause lookup from unit_update 'deactivated' event notes."""
    try:
        events_df = _er.get_events(
            event_category="monitoring",
            event_type=[UNIT_UPDATE_EVENT_TYPE],
            include_details=True,
            drop_null_geometry=False,
        )
    except Exception:
        return {}
    if events_df is None or events_df.empty:
        return {}
    latest = {}
    for _, ev in events_df.iterrows():
        details = ev.get("event_details") or {}
        if isinstance(details, str):
            try:
                details = json.loads(details)
            except Exception:
                details = {}
        if not isinstance(details, dict):
            continue
        unit_id = details.get("unitupdate_unitid")
        action = str(details.get("unitupdate_action", "")).lower()
        if unit_id in source_ids_tuple and "deactiv" in action:
            notes = (details.get("unitupdate_notes") or "").strip()
            time_val = ev.get("time")
            prev = latest.get(unit_id)
            if notes and (prev is None or (time_val and str(time_val) > str(prev[1]))):
                latest[unit_id] = (notes, time_val)
    return {k: v[0] for k, v in latest.items()}


# ═══════════════════════════════════════════════════════════════════════════
# Scheduled-fix detection
# ═══════════════════════════════════════════════════════════════════════════

def _auto_scheduled_fixes(mode_daily, is_gsat):
    if is_gsat:
        if mode_daily >= 40:
            return 48
        elif mode_daily >= 20:
            return 24
        elif mode_daily >= 8:
            return 12
        elif mode_daily >= 3:
            return 4
        return 4
    else:
        if mode_daily >= 90:
            return 96
        elif mode_daily >= 40:
            return 48
        elif mode_daily >= 20:
            return 24
        elif mode_daily >= 10:
            return 12
        elif mode_daily >= 6:
            return 8
        elif mode_daily >= 3:
            return 4
        return 24


def _mode(series):
    counts = series.value_counts()
    return counts.index[0] if len(counts) else np.nan


# ═══════════════════════════════════════════════════════════════════════════
# Core analysis
# ═══════════════════════════════════════════════════════════════════════════

def analyze_manufacturer(label, cfg, sources_df, assignments_df, obs_df, deactivation_notes,
                          country_region_map, subject_metadata):
    """Compute every stat needed for one device-type section of the report."""
    now = pd.Timestamp.now(tz="UTC")
    is_gsat = cfg["provider"] == "gsatsolar"
    is_spoortrack = cfg["provider"] == "spoortrack"

    src = sources_df.copy()
    if cfg.get("exclude_tagid_contains"):
        src = src[~src["tagID"].str.contains(cfg["exclude_tagid_contains"], case=False, na=False)]

    asn = assignments_df[assignments_df["source_id"].isin(src["id"].astype(str))].copy()
    if cfg.get("since_filter") is not None:
        cutoff = pd.Timestamp(cfg["since_filter"], tz="UTC")
        asn = asn[asn["depStart"] >= cutoff]
    if asn.empty:
        return None

    # Restrict to giraffe subjects only — some of these device types are also
    # deployed on other species for unrelated research and must not leak in here.
    # Strict equality (not "anything not explicitly excluded"): unknown/missing
    # subtype is excluded too, matching the R script's filter(subType=="giraffe").
    asn["subtype"] = asn["subject_id"].map(lambda s: (subject_metadata.get(s) or {}).get("subtype", ""))
    asn = asn[asn["subtype"] == GIRAFFE_SUBJECT_SUBTYPE].drop(columns=["subtype"])
    if asn.empty:
        return None

    # Authoritative subject name from subjects/ — the subjectsources relation doesn't
    # reliably expand to a named object, so prefer this lookup over whatever came
    # through on the assignment row itself.
    resolved_name = asn["subject_id"].map(lambda s: (subject_metadata.get(s) or {}).get("name"))
    asn["subName"] = resolved_name.where(resolved_name.notna() & (resolved_name != ""), asn["subName"])

    asn["depEndEffective"] = asn["depEnd"].fillna(now)
    asn.loc[asn["depEndEffective"] > now, "depEndEffective"] = now
    asn["status"] = np.where(asn["depEnd"].isna() | (asn["depEnd"] >= now), "Active", "Ended")
    asn["deployment_length_days"] = (asn["depEndEffective"] - asn["depStart"]).dt.days
    asn["end_cause"] = asn["source_id"].map(lambda s: deactivation_notes.get(s, ""))
    asn["end_cause"] = np.where(asn["status"] == "Active", "", asn["end_cause"])
    asn.loc[(asn["status"] != "Active") & (asn["end_cause"] == ""), "end_cause"] = "[No cause recorded]"
    _group_default = (None, None, None, None)
    asn["country"] = asn["subject_id"].map(lambda s: (country_region_map.get(s) or _group_default)[0])
    asn["region"] = asn["subject_id"].map(lambda s: (country_region_map.get(s) or _group_default)[1])
    asn["group_name"] = asn["subject_id"].map(lambda s: (country_region_map.get(s) or _group_default)[2])

    tag_map = dict(zip(src["id"].astype(str), src["tagID"]))
    asn["tagID"] = asn["source_id"].map(tag_map)

    # ── Per-source deployment window (for first/last-day exclusion) ──────────
    dep_window = asn.groupby("source_id").agg(dep_start=("depStart", "min"), dep_end=("depEndEffective", "max"))

    # ── Trim the first & last calendar day of each unit's deployment ─────────
    # Activation/deactivation happens partway through those days, so both
    # location counts and battery readings are artificially low on them —
    # exclude from every battery/fix-rate metric and chart (matches the R report's
    # existing fix-rate handling, now applied consistently to battery too).
    obs_trimmed = obs_df.copy()
    obs_trimmed["date"] = obs_trimmed["obsDatetime"].dt.date
    obs_trimmed = obs_trimmed.merge(dep_window, left_on="source_id", right_index=True, how="left")
    obs_trimmed["dep_start_date"] = obs_trimmed["dep_start"].dt.date
    obs_trimmed["dep_end_date"] = obs_trimmed["dep_end"].dt.date
    obs_trimmed = obs_trimmed[
        (obs_trimmed["date"] > obs_trimmed["dep_start_date"]) & (obs_trimmed["date"] < obs_trimmed["dep_end_date"])
    ].drop(columns=["dep_start", "dep_end", "dep_start_date", "dep_end_date"])

    # ── Battery (trimmed) ────────────────────────────────────────────────────
    obs_batt = obs_trimmed.dropna(subset=["battery"]).copy()
    obs_batt = obs_batt[obs_batt["battery"] > 0]
    obs_batt = obs_batt.sort_values(["source_id", "obsDatetime"])

    battery_rows, cv_rows = [], []
    for sid, g in obs_batt.groupby("source_id"):
        if g.empty:
            continue
        start_v = g["battery"].iloc[0]
        current_v = g["battery"].iloc[-1]
        days = (g["obsDatetime"].iloc[-1] - g["obsDatetime"].iloc[0]).days
        drop = start_v - current_v
        rate = drop / days if days > 0 else 0.0
        floor = cfg["battery_floor"]
        projected = (current_v - floor) / rate if rate > 0 else np.nan
        battery_rows.append(dict(
            source_id=sid, start_voltage=start_v, current_voltage=current_v,
            days_deployed=days, voltage_drop_per_day=rate, projected_life_days=projected,
        ))
        if len(g) >= 5:
            m, sdv = g["battery"].mean(), g["battery"].std()
            cv_rows.append(dict(source_id=sid, cv=(sdv / m if m else np.nan)))
    battery_df = pd.DataFrame(battery_rows) if battery_rows else pd.DataFrame(
        columns=["source_id", "start_voltage", "current_voltage", "days_deployed",
                 "voltage_drop_per_day", "projected_life_days"]
    )
    cv_df = pd.DataFrame(cv_rows) if cv_rows else pd.DataFrame(columns=["source_id", "cv"])

    # ── Fix rate (trimmed) — counts every location fix, not just battery-bearing ones ──
    daily = obs_trimmed.groupby(["source_id", "date"]).size().reset_index(name="daily_locations")

    schedule_rows = []
    for sid, g in daily.groupby("source_id"):
        mode_daily = _mode(g["daily_locations"])
        schedule_rows.append(dict(source_id=sid, mode_daily=mode_daily,
                                   scheduled_fixes=_auto_scheduled_fixes(mode_daily, is_gsat)))
    schedule_df = pd.DataFrame(schedule_rows)
    daily = daily.merge(schedule_df[["source_id", "scheduled_fixes"]], on="source_id", how="left")

    tag_to_sources = {}
    for s, t in tag_map.items():
        tag_to_sources.setdefault(t, []).append(s)

    if is_gsat:
        for sid_set, fixes in [(GSAT_MANUAL_6PERDAY, 6), (GSAT_MANUAL_12PERDAY, 12)]:
            mask = daily["source_id"].map(lambda s: tag_map.get(s) in sid_set)
            daily.loc[mask, "scheduled_fixes"] = fixes
        for tag_id in GSAT_SCHEDULE_CHANGE_UNITS:
            for sid in tag_to_sources.get(tag_id, []):
                mask = (daily["source_id"] == sid)
                daily.loc[mask & (daily["date"] < GSAT_SCHEDULE_CHANGE_DATE), "scheduled_fixes"] = 12
                daily.loc[mask & (daily["date"] >= GSAT_SCHEDULE_CHANGE_DATE), "scheduled_fixes"] = 24
    if is_spoortrack:
        for tag_id, periods in SPOORTRACK_TEST_PERIODS.items():
            for sid in tag_to_sources.get(tag_id, []):
                mask = (daily["source_id"] == sid)
                for start_p, end_p, fixes in periods:
                    period_mask = mask & (daily["date"] >= start_p) & (daily["date"] <= end_p)
                    daily.loc[period_mask, "scheduled_fixes"] = fixes

    daily["fix_rate"] = daily["daily_locations"] / daily["scheduled_fixes"]
    overall_mean_fix_rate = daily["fix_rate"].mean() if not daily.empty else np.nan
    per_source_fix_rate = daily.groupby("source_id")["fix_rate"].mean().rename("mean_fix_rate")

    # ── Summary table (one row per assignment) ──────────────────────────────
    # last_transmission/days_since_last_tx use the FULL (untrimmed) history —
    # this is a liveness check ("is the unit still talking to us"), not a
    # performance metric, so it shouldn't discard the most recent day's data.
    last_tx = obs_df.groupby("source_id")["obsDatetime"].max().rename("last_transmission")
    summary = asn.merge(last_tx, on="source_id", how="left")
    summary["days_since_last_tx"] = (now - summary["last_transmission"]).dt.days
    summary = summary.merge(
        battery_df[["source_id", "current_voltage"]].rename(columns={"current_voltage": "latest_battery"}),
        on="source_id", how="left",
    )
    summary = summary.sort_values("deployment_length_days", ascending=False).reset_index(drop=True)

    # ── Country / region breakdown ──────────────────────────────────────────
    loc = summary.dropna(subset=["country"]).groupby(["country", "region"]).agg(
        n_units=("source_id", "count"), n_active=("status", lambda s: (s == "Active").sum())
    ).reset_index().sort_values(["country", "n_units"], ascending=[True, False])
    country_totals = loc.groupby("country").agg(total_units=("n_units", "sum"), total_active=("n_active", "sum")) \
        .reset_index().sort_values("total_units", ascending=False)

    # ── Performance scoring ──────────────────────────────────────────────────
    perf = summary.merge(per_source_fix_rate, on="source_id", how="left")
    perf = perf.merge(cv_df, on="source_id", how="left")
    perf["fix_score"] = perf["mean_fix_rate"].clip(lower=0, upper=1)
    perf["battery_score"] = (1 - (perf["cv"] / 0.5)).clip(lower=0, upper=1)
    perf["longevity_score"] = (perf["deployment_length_days"] / 365).clip(lower=0, upper=1)
    score_cols = ["fix_score", "battery_score", "longevity_score"]
    all_present = perf[score_cols].notna().all(axis=1)
    perf["overall_score"] = np.where(all_present, perf[score_cols].mean(axis=1), np.nan)
    scored = perf.dropna(subset=["overall_score"])
    total_scored = len(scored)
    excellent = int((scored["overall_score"] > 0.8).sum())
    good = int(((scored["overall_score"] >= 0.6) & (scored["overall_score"] <= 0.8)).sum())
    poor = int((scored["overall_score"] < 0.6).sum())

    first_deployed = summary["depStart"].min()

    return dict(
        label=label, cfg=cfg, summary=summary, loc=loc, country_totals=country_totals,
        battery_df=battery_df, cv_df=cv_df, battery_obs=obs_batt, daily=daily,
        overall_mean_fix_rate=overall_mean_fix_rate,
        excellent=excellent, good=good, poor=poor, total_scored=total_scored,
        first_deployed=first_deployed,
    )


# ═══════════════════════════════════════════════════════════════════════════
# Charts (matplotlib, returned as PNG bytes for both st display & docx embed)
# ═══════════════════════════════════════════════════════════════════════════

def _fig_to_png(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def chart_deployment(result):
    df = result["summary"].copy()
    df["loc_label"] = df["country"].fillna("Unknown") + np.where(
        df["region"].notna(), " - " + df["region"].fillna(""), ""
    )
    fig, ax = plt.subplots(figsize=(9, 5))
    categories = sorted(df["loc_label"].unique())
    cat_x = {c: i for i, c in enumerate(categories)}
    rng = random.Random(42)
    colors = {"Active": "#2E8B57", "Ended": "#CD5C5C"}
    for status, group in df.groupby("status"):
        xs = [cat_x[c] + rng.uniform(-0.2, 0.2) for c in group["loc_label"]]
        ax.scatter(xs, group["deployment_length_days"], s=40, alpha=0.8,
                   color=colors.get(status, "gray"), label=status)
    mean_len = df["deployment_length_days"].mean()
    ax.axhline(mean_len, linestyle="--", color="black", alpha=0.7)
    ax.text(len(categories) - 0.5, mean_len, f"Mean: {mean_len:.0f} days", va="bottom", ha="right", fontsize=9)
    ax.set_xticks(range(len(categories)))
    ax.set_xticklabels(categories, rotation=45, ha="right")
    ax.set_ylabel("Deployment length (days)")
    ax.set_title(f"{result['label']} — deployment duration by location", loc="right")
    ax.legend(loc="upper left", bbox_to_anchor=(1.0, 1.0))
    fig.tight_layout()
    return _fig_to_png(fig)


def chart_battery(result):
    unit = result["cfg"]["battery_unit"]
    obs = result["battery_obs"]
    daily_mean = obs.groupby(["source_id", "date"])["battery"].mean().reset_index()
    fig, ax = plt.subplots(figsize=(9, 5))
    for sid, g in daily_mean.groupby("source_id"):
        ax.plot(g["date"], g["battery"], alpha=0.5, linewidth=1)
    overall_mean = daily_mean["battery"].mean()
    ax.axhline(overall_mean, linestyle="--", color="black", linewidth=1.5, alpha=0.8)
    label = f"Mean: {overall_mean:.2f} {'V' if unit == 'voltage' else '%'}"
    ax.text(daily_mean["date"].max(), overall_mean, label, ha="right", va="bottom", fontsize=9)
    ax.set_ylabel("Battery voltage (V)" if unit == "voltage" else "Battery percentage (%)")
    ax.set_xlabel("Date")
    ax.set_title(f"{result['label']} — battery performance over time", loc="right")
    fig.autofmt_xdate(rotation=45)
    fig.tight_layout()
    return _fig_to_png(fig)


def chart_fixrate(result):
    daily = result["daily"]
    fig, ax = plt.subplots(figsize=(9, 5))
    for sid, g in daily.groupby("source_id"):
        ax.plot(g["date"], g["fix_rate"], alpha=0.5, linewidth=1)
    overall_mean = daily["fix_rate"].mean() if not daily.empty else 0
    ax.axhline(overall_mean, linestyle="--", color="black", linewidth=1.5, alpha=0.8)
    ax.text(daily["date"].max() if not daily.empty else 0, overall_mean,
            f"Mean: {overall_mean:.2f}", ha="right", va="bottom", fontsize=9)
    ax.set_ylabel("Fix rate (received / scheduled)")
    ax.set_xlabel("Date")
    ax.set_title(f"{result['label']} — fix rate performance (actual/scheduled)", loc="right")
    fig.autofmt_xdate(rotation=45)
    fig.tight_layout()
    return _fig_to_png(fig)


# ═══════════════════════════════════════════════════════════════════════════
# DOCX report building
# ═══════════════════════════════════════════════════════════════════════════

def _add_toc(document):
    """Insert a real Word TOC field (Google Docs will offer to update/convert it on open)."""
    try:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-2" \\h \\z \\u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "Right-click and choose 'Update Field' to build the table of contents."
        fld_sep.append(placeholder)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r = run._r
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(fld_end)
    except Exception:
        document.add_paragraph(
            "(Insert > Table of contents in Google Docs once this file is opened, "
            "using the headings below.)"
        )


def build_docx(author, report_date, results, comments_by_label):
    document = Document()

    if GCF_LOGO.exists():
        try:
            document.add_picture(str(GCF_LOGO), width=Inches(1.6))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    title = document.add_heading("GPS unit performance", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author)
    run.italic = True
    p2 = document.add_paragraph(f"Report date: {report_date}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    document.add_paragraph("GPS unit performance for tags deployed on giraffe species.")
    document.add_paragraph()

    _add_toc(document)
    document.add_paragraph()

    year = report_date.year if hasattr(report_date, "year") else datetime.now().year
    document.add_paragraph(
        f"Recommended citation: {author.split(',')[0].strip()} ({year}) "
        f"GPS unit performance report {report_date}. Giraffe Conservation Foundation, Windhoek, Namibia."
    )
    document.add_page_break()

    for result in results:
        label = result["label"]
        cfg = result["cfg"]
        summary = result["summary"]
        unit_suffix = "V" if cfg["battery_unit"] == "voltage" else "%"

        document.add_heading(label, level=1)

        # ── Deployments ──────────────────────────────────────────────────
        document.add_heading("Deployments", level=2)
        document.add_paragraph(f"Unit type first deployed: {result['first_deployed'].date()}")
        document.add_paragraph(f"Total units deployed: {len(summary)}")
        document.add_paragraph(f"Currently active: {(summary['status'] == 'Active').sum()}")
        document.add_paragraph(f"Ended deployments: {(summary['status'] != 'Active').sum()}")
        document.add_paragraph()

        for _, row in result["country_totals"].iterrows():
            document.add_paragraph(f"{row['country']}: {row['total_units']} units ({row['total_active']} active)")
            regions = result["loc"][result["loc"]["country"] == row["country"]]
            for _, rrow in regions.iterrows():
                if rrow["region"]:
                    document.add_paragraph(
                        f"{rrow['region']}: {rrow['n_units']} units ({rrow['n_active']} active)", style="List Bullet"
                    )
            document.add_paragraph()

        mean_len = summary["deployment_length_days"].mean()
        lengths_desc = summary["deployment_length_days"].sort_values(ascending=False)
        lengths_asc = summary["deployment_length_days"].sort_values(ascending=True)
        document.add_paragraph(f"Mean deployment length: {mean_len:.0f} days")
        if len(lengths_desc) >= 3:
            document.add_paragraph(
                "Top 3 longest deployments: " + ", ".join(f"{d} days" for d in lengths_desc.head(3))
            )
            document.add_paragraph(
                "Bottom 3 shortest deployments: " + ", ".join(f"{d} days" for d in lengths_asc.head(3))
            )
        not_transmitting = summary[(summary["status"] == "Active") & (summary["days_since_last_tx"] > 30)]
        document.add_paragraph(f"Active units not transmitting >30 days: {len(not_transmitting)}")
        document.add_picture(io.BytesIO(chart_deployment(result)), width=Inches(6))
        document.add_paragraph(f"Figure 1. Deployment lengths of {label} units deployed since "
                                f"{result['first_deployed'].strftime('%d %B %Y')}.")
        document.add_page_break()

        # ── Battery ──────────────────────────────────────────────────────
        document.add_heading("Battery", level=2)
        mean_batt = summary["latest_battery"].mean()
        rng = cfg["battery_range"]
        document.add_paragraph(
            f"Mean current battery {'voltage' if cfg['battery_unit'] == 'voltage' else 'percentage'}: "
            f"{mean_batt:.2f}{unit_suffix} (functional range = {rng[0]} - {rng[1]}{unit_suffix})"
        )
        under_30 = int((result["battery_df"]["projected_life_days"] < 30).sum()) if not result["battery_df"].empty else 0
        document.add_paragraph(f"Units with projected battery life <30 days: {under_30}")
        mean_cv = result["cv_df"]["cv"].mean() if not result["cv_df"].empty else float("nan")
        document.add_paragraph(
            f"Mean coefficient of variation (battery consistency, lower=more consistent): {mean_cv:.2f}"
        )
        consistent = int((result["cv_df"]["cv"] < 0.3).sum()) if not result["cv_df"].empty else 0
        document.add_paragraph(f"Highly consistent batteries (CV < 0.3): {consistent} / {len(result['cv_df'])}")
        document.add_picture(io.BytesIO(chart_battery(result)), width=Inches(6))
        document.add_paragraph(f"Figure 2. Mean battery {cfg['battery_unit']} of {label} units deployed since "
                                f"{result['first_deployed'].strftime('%d %B %Y')}.")
        document.add_page_break()

        # ── Success rate ───────────────────────────────────────────────────
        document.add_heading("Success rate", level=2)
        document.add_paragraph(f"Mean fix rate: {result['overall_mean_fix_rate']:.2f} (recorded/scheduled)")
        document.add_picture(io.BytesIO(chart_fixrate(result)), width=Inches(6))
        document.add_paragraph(f"Figure 3. Fix rate (received/scheduled) of {label} units deployed since "
                                f"{result['first_deployed'].strftime('%d %B %Y')}.")
        # Single combined view: grouped by site (full history) so regional clustering
        # is visible, with raw individual cause text (not bucketed into categories)
        # under each site — no separate flat aggregate list alongside it.
        ended = summary[summary["status"] != "Active"].copy()
        if not ended.empty:
            ended["site_label"] = ended["group_name"].fillna("unknown group").replace("", "unknown group")

            document.add_paragraph("Failure causes:")
            for site, total in ended["site_label"].value_counts().items():
                site_p = document.add_paragraph()
                site_run = site_p.add_run(f"{site} — {total} total")
                site_run.bold = True
                site_ended = ended[ended["site_label"] == site]
                for cause, n in site_ended["end_cause"].value_counts().items():
                    document.add_paragraph(f"{cause}: {n}", style="List Bullet")
        document.add_page_break()

        # ── Overall assessment ─────────────────────────────────────────────
        document.add_heading("Overall assessment", level=2)
        document.add_paragraph(
            "Performance distribution (overall score = fix success + battery reliability + longevity):"
        )
        total_scored = max(result["total_scored"], 1)
        document.add_paragraph(
            f"Excellent (>0.8): {result['excellent']} units ({100 * result['excellent'] / total_scored:.0f}%)",
            style="List Bullet",
        )
        document.add_paragraph(
            f"Good (0.6-0.8): {result['good']} units ({100 * result['good'] / total_scored:.0f}%)",
            style="List Bullet",
        )
        document.add_paragraph(
            f"Poor (<0.6): {result['poor']} units ({100 * result['poor'] / total_scored:.0f}%)",
            style="List Bullet",
        )
        document.add_paragraph()
        document.add_paragraph("General comments:")
        comment_text = comments_by_label.get(label, "").strip()
        document.add_paragraph(comment_text if comment_text else "(add general comments here)")
        document.add_paragraph()

        excellent_pct = 100 * result["excellent"] / total_scored
        good_pct = 100 * result["good"] / total_scored
        if excellent_pct > 70:
            recommendation = "CONTINUE - strong performance"
        elif excellent_pct + good_pct > 80:
            recommendation = "CONTINUE with monitoring - mixed performance"
        else:
            recommendation = "CONSIDER ALTERNATIVES - poor overall performance"
        p_final = document.add_paragraph()
        run = p_final.add_run(f"OVERALL ASSESSMENT: {recommendation.upper()}")
        run.bold = True
        document.add_page_break()

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# Streamlit UI
# ═══════════════════════════════════════════════════════════════════════════

def _main_implementation():
    st.title("📡 Unit Performance Report")
    st.caption(
        "Generates the quarterly GPS unit performance report (deployments, battery, fix rate, "
        "overall assessment) for SpoorTrack, Savannah Tracking and GSat Solar tags, formatted as "
        "a downloadable Word document you can upload to Drive and refine further."
    )

    er = require_earthranger_login("Unit Performance Report")

    st.info(
        "This pulls **full device history** from EarthRanger for every unit ever deployed on each "
        "programme, so it can take a few minutes for large fleets. Results are cached for 30–60 minutes."
    )

    with st.expander("⚙️ Report settings", expanded=True):
        author = st.text_input("Author / citation name", value="Courtney Marneweck, GCF")
        report_date = st.date_input("Report date", value=date.today())
        st.markdown("**General comments** (optional — appears in the Overall Assessment of each section)")
        comments_by_label = {}
        for label in MANUFACTURERS:
            comments_by_label[label] = st.text_area(f"{label} comments", key=f"comments_{label}", height=90)

        st.markdown("---")
        override_range = st.checkbox(
            "Override observation date range (for faster testing or a custom-period report)",
            value=False,
        )
        custom_since, custom_until = None, None
        if override_range:
            c1, c2 = st.columns(2)
            custom_since = c1.date_input("Observations from", value=date.today() - pd.Timedelta(days=14))
            custom_until = c2.date_input("Observations to", value=date.today())
            st.caption(
                "⚠️ This only narrows the (slow) observation-history fetch used for battery, "
                "fix-rate, and their charts — deployment counts, status, and duration stay "
                "full-history regardless, since those come from a separate, cheap lookup. "
                "A narrow window here is much faster but won't match the cumulative-since-inception "
                "framing of a normal quarterly report — use this for quick testing or a "
                "deliberately scoped custom report, not the standard quarterly output."
            )

    if not st.button("🚀 Generate report", type="primary"):
        return

    results = []
    with st.spinner("Fetching subject groups (country/region)..."):
        country_region_map, group_debug_info = fetch_subject_country_region(er)
    with st.spinner("Fetching subject names and species..."):
        subject_metadata = fetch_subject_metadata(er)
    if not country_region_map:
        st.warning(
            "Could not resolve any subject groups (country/region/site will show as unknown). "
            f"Debug info: {group_debug_info}"
        )
    if not subject_metadata:
        st.error(
            "Could not resolve any subject names/species from EarthRanger — the giraffe-only "
            "filter can't be applied, so the report would be empty. Check EarthRanger "
            "permissions for the subjects/ endpoint and try again."
        )
        return

    for label, cfg in MANUFACTURERS.items():
        with st.spinner(f"Fetching {label} sources..."):
            sources_df = fetch_manufacturer_sources(er, cfg["provider"])
        if sources_df.empty:
            st.warning(f"No {label} tracking sources found in EarthRanger — skipping.")
            continue

        source_ids = tuple(sorted(sources_df["id"].astype(str).unique()))

        with st.spinner(f"Fetching {label} deployment history ({len(source_ids)} units)..."):
            assignments_df = fetch_assignments(er, source_ids)
        if assignments_df.empty:
            st.warning(f"No deployment history found for {label} — skipping.")
            continue

        if override_range:
            since_iso = pd.Timestamp(custom_since, tz="UTC").isoformat()
            until_iso = (pd.Timestamp(custom_until, tz="UTC") + pd.Timedelta(days=1)).isoformat()
        else:
            since_iso = (
                pd.Timestamp(cfg["since_filter"], tz="UTC").isoformat()
                if cfg.get("since_filter")
                else assignments_df["depStart"].min().isoformat()
            )
            until_iso = pd.Timestamp.now(tz="UTC").isoformat()

        with st.spinner(f"Fetching {label} observation history — this is the slow part..."):
            obs_df = fetch_observation_history(er, source_ids, since_iso, until_iso, cfg["battery_unit"])

        with st.spinner(f"Checking {label} deactivation notes..."):
            deactivation_notes = fetch_deactivation_notes(er, source_ids)

        result = analyze_manufacturer(label, cfg, sources_df, assignments_df, obs_df, deactivation_notes,
                                       country_region_map, subject_metadata)
        if result is None:
            st.warning(f"No usable data for {label} — skipping.")
            continue
        results.append(result)

    if not results:
        st.error("No data available for any device type — nothing to report.")
        return

    st.markdown("---")
    st.subheader("📋 Preview")
    tabs = st.tabs([r["label"] for r in results])
    for tab, result in zip(tabs, results):
        with tab:
            c1, c2, c3, c4 = st.columns(4)
            summary = result["summary"]
            c1.metric("Total units", len(summary))
            c2.metric("Active", int((summary["status"] == "Active").sum()))
            c3.metric("Mean fix rate", f"{result['overall_mean_fix_rate']:.2f}")
            c4.metric("Mean deployment (days)", f"{summary['deployment_length_days'].mean():.0f}")
            st.image(chart_deployment(result), use_container_width=True)
            st.image(chart_battery(result), use_container_width=True)
            st.image(chart_fixrate(result), use_container_width=True)

    with st.spinner("Building Word document..."):
        docx_bytes = build_docx(author, report_date, results, comments_by_label)

    st.markdown("---")
    st.success("Report generated.")
    st.download_button(
        "⬇️ Download report (.docx)",
        data=docx_bytes,
        file_name=f"GCF_unitPerformance_{report_date.strftime('%y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
    )
    st.caption(
        "Upload this to Drive and open with Google Docs (or File → Open with Google Docs). "
        "Update the table of contents field and adjust wording as needed."
    )


if __name__ == "__main__":
    main()
